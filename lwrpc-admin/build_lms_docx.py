import os
import zipfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image, ImageFilter

# ---------------------------------------------------------------------------
# COLOR PALETTE TOKENS
# ---------------------------------------------------------------------------
COLOR_NAVY = RGBColor(10, 37, 64)       # #0A2540
COLOR_ROYAL = RGBColor(29, 78, 216)     # #1D4ED8
COLOR_EMERALD = RGBColor(5, 150, 105)   # #059669
COLOR_CRIMSON = RGBColor(220, 38, 38)   # #DC2626
COLOR_TEXT = RGBColor(30, 41, 59)       # #1E293B
HEX_NAVY = "0A2540"
HEX_LIGHT_BG = "F8FAFC"
HEX_CALLOUT_BORDER = "1D4ED8"

# ---------------------------------------------------------------------------
# HELPER FUNCTIONS FOR WORD XML FORMATTING
# ---------------------------------------------------------------------------
def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, text, title="OPERATIONAL RULE / WARNING"):
    """Adds a styled callout box with a colored left accent border."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border formatting
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{HEX_CALLOUT_BORDER}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    r_title = p.add_run(f"[{title}]\n")
    r_title.bold = True
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = COLOR_ROYAL
    
    r_body = p.add_run(text)
    r_body.font.size = Pt(9.5)
    r_body.font.color.rgb = COLOR_TEXT
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_heading_1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    for r in h.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = COLOR_NAVY
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    for r in h.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = COLOR_ROYAL
    return h

def insert_screenshot_if_exists(doc, image_path, caption_text, width=Inches(6.0)):
    """Inserts a screenshot with caption if the image file exists."""
    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        run = p_img.add_run()
        run.add_picture(image_path, width=width)
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run(f"Figure: {caption_text}")
        r_cap.font.italic = True
        r_cap.font.size = Pt(9)
        r_cap.font.color.rgb = RGBColor(100, 116, 139)
    else:
        add_callout_box(doc, f"Screenshot File Missing: {image_path}. Place image in target directory.", title="IMAGE PLACEHOLDER")

# ---------------------------------------------------------------------------
# DOCUMENT GENERATION
# ---------------------------------------------------------------------------
def build_lms_document():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Header / Footer setup
        footer = section.footer
        p_ft = footer.paragraphs[0]
        p_ft.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_ft = p_ft.add_run("Lakewood Ranch Pickleball Club | LMS Master User Guide v0566")
        r_ft.font.size = Pt(8.5)
        r_ft.font.color.rgb = RGBColor(148, 163, 184)

    # -----------------------------------------------------------------------
    # COVER PAGE
    # -----------------------------------------------------------------------
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists("LWRPC-New.jpg"):
        p_logo.add_run().add_picture("LWRPC-New.jpg", width=Inches(2.2))
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    r = p_title.add_run("LAKEWOOD RANCH PICKLEBALL CLUB\n")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = COLOR_ROYAL
    
    r_sub = p_title.add_run("LEAGUE MANAGEMENT SYSTEM (LMS)\n")
    r_sub.font.size = Pt(24)
    r_sub.font.bold = True
    r_sub.font.color.rgb = COLOR_NAVY
    
    r_man = p_title.add_run("Master Operations Manual & System Reference")
    r_man.font.size = Pt(16)
    r_man.font.color.rgb = COLOR_EMERALD

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(150)
    r_m = p_meta.add_run("System Build: LMS-0566 | Platform: Next.js 16 / Supabase\nDate: Fall 2026 Season Release\nAuthor: League Administration & Operations")
    r_m.font.size = Pt(10)
    r_m.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # TABLE OF CONTENTS PLACEHOLDER
    # -----------------------------------------------------------------------
    add_heading_1(doc, "Table of Contents")
    p_toc = doc.add_paragraph()
    run_toc = p_toc.add_run("[Table of Contents field - Right-click in Word and select 'Update Field']")
    run_toc.font.italic = True
    run_toc.font.color.rgb = COLOR_ROYAL
    doc.add_page_break()

    # -----------------------------------------------------------------------
    # SECTION 1: SYSTEM OVERVIEW
    # -----------------------------------------------------------------------
    add_heading_1(doc, "1. System Overview & Architecture")
    p = doc.add_paragraph("The Lakewood Ranch Pickleball Club (LWRPC) League Management System (LMS) is a complete administrative platform designed specifically for club leagues, ratings, match operations, schedules, and standings.")
    p.runs[0].font.size = Pt(10.5)

    insert_screenshot_if_exists(doc, "01-admin-dashboard-01-overview.png", "Admin Dashboard Overview & Operational Metrics Center")

    add_callout_box(doc, "The LMS separates causal social play from official DUPR competition. Standalone round robins run on separate database tables to protect league standings.", title="ARCHITECTURE ARCHITECTURE NOTE")

    # -----------------------------------------------------------------------
    # SECTION 2: PLAYER & CAPTAIN MANUALS
    # -----------------------------------------------------------------------
    add_heading_1(doc, "2. Player & Team Captain Manual")
    
    add_heading_2(doc, "2.1 Player Dashboard")
    doc.add_paragraph("Players can log in at https://league.lwrpickleballclub.com to view upcoming match locations, court numbers, team standings, and play history.")
    
    insert_screenshot_if_exists(doc, "02-members-01.png", "Player Dashboard and Team Roster Controls")

    add_heading_2(doc, "2.2 Team Captain Match Setup & Lineups")
    doc.add_paragraph("Captains must exchange lineups via Match Setup at least 3 days before scheduled play. The system checks individual ratings and team maximum caps automatically.")

    insert_screenshot_if_exists(doc, "20-score-sheets-01.png", "Match Setup & Printable Score Sheet Format")

    add_callout_box(doc, "If a match stops before 6 points, record it as a Forfeit (0-0, not sent to DUPR). If 6+ points were played, record as Retired (scores stand and post to DUPR).", title="FORFEIT VS. RETIREMENT RULE")

    # -----------------------------------------------------------------------
    # SECTION 3: SYSTEM ADMINISTRATOR MANUAL
    # -----------------------------------------------------------------------
    add_heading_1(doc, "3. System Administrator Manual")

    add_heading_2(doc, "3.1 Visual Schedule Editor & Overbooking Engine")
    doc.add_paragraph("League Managers can generate automated schedules, resolve court overbooking conflicts, and swap home/away teams with instant capacity analysis.")

    insert_screenshot_if_exists(doc, "19-round-robin-manager-01.png", "Visual Schedule Editor and Match Controls")

    add_heading_2(doc, "3.2 Scoring Operations & DUPR Export")
    doc.add_paragraph("Scoring operations allow administrators to audit score entries, override disputed matches, and generate official DUPR CSV export batches.")

    insert_screenshot_if_exists(doc, "23-member-import-01.png", "Scoring Operations & Member Import Tools")

    # Save output
    output_filename = "LWRPC_LMS_Master_User_Guide.docx"
    doc.save(output_filename)
    print(f"Successfully generated master user guide: {output_filename}")

if __name__ == "__main__":
    build_lms_document()